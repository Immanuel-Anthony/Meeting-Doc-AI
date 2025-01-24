import os
import subprocess
from faster_whisper import WhisperModel
import speech_recognition as sr
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime
import openai  
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import csv

def extract_emails(csv_path):
    emails = []
    with open(csv_path, 'r') as file:
        reader = csv.DictReader(file)
        for row in reader:
            if 'Email' in row:
                emails.append(row['Email'])
    return emails

def send_email(sender_email, sender_password, recipient_email, subject, body, attachment_path):
    try:
        # Create the email
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = subject

        # Attach the email body
        msg.attach(MIMEText(body, 'plain'))

        # Attach the .docx file
        with open(attachment_path, 'rb') as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename={attachment_path.split("/")[-1]}'
            )
            msg.attach(part)

        # Send the email
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())
        print(f"Email sent to {recipient_email}")
    except Exception as e:
        print(f"Failed to send email to {recipient_email}: {e}")

class MeetingMinutesGenerator:
    def __init__(self, model_size="base", device="cpu", compute_type="float32"):
        # Load OpenAI API Key from environment variable
        openai.api_key = os.getenv('OPENAI_API_KEY')
        
        # Audio and Whisper model parameters
        self.model_size = model_size
        self.device = device
        self.compute_type = compute_type
        self.model = WhisperModel(model_size, device=device, compute_type=compute_type)

    def extract_audio(self, video_file):
        # Add detailed error checking
        if not os.path.exists(video_file):
            raise FileNotFoundError(f"Video file {video_file} not found.")
        # Check if ffmpeg is installed
        try:
            subprocess.run(["ffmpeg", "-version"], capture_output=True, text=True)
        except FileNotFoundError:
            raise FileNotFoundError("FFmpeg is not installed or not in system PATH")

        audio_file = "output.wav"
        command = ["ffmpeg", "-i", video_file, "-vn", "-ac", "1", "-ar", "16000", "-acodec", "pcm_s16le", audio_file]
        
        try:
            result = subprocess.run(command, capture_output=True, text=True, check=True)
            print("Audio extraction completed.")
            return audio_file
        except subprocess.CalledProcessError as e:
            print(f"FFmpeg Error: {e}")
            print("STDOUT:", e.stdout)
            print("STDERR:", e.stderr)
            raise

    def transcribe_audio(self, audio_file):
        """Transcribe audio using Whisper"""
        if not os.path.exists(audio_file):
            raise FileNotFoundError(f"Audio file {audio_file} not found.")
        
        print("Transcribing audio...")
        try:
            segments, info = self.model.transcribe(audio_file, beam_size=5)
            transcription = " ".join([segment.text for segment in segments])
            return transcription
        except Exception as e:
            print(f"Transcription error: {e}")
            raise
   
    def summarize_transcription(self, transcription):
        if not openai.api_key:
            raise ValueError("OpenAI API key is required for dynamic summarization")
        
        try:
            client = openai.OpenAI()
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Provide a concise 2-3 sentence executive summary of the meeting, capturing the most critical insights and overall purpose of the discussion. Then generate detailed meeting minutes with Key Discussion Points, Action Items, and Additional Notes."},
                    {"role": "user", "content": transcription}
                ],
                max_tokens=400
            )
            full_response = response.choices[0].message.content
            
            # Split the response into executive summary and rest of the minutes
            parts = full_response.split("\n\n", 1)
            executive_summary = parts[0]
            detailed_minutes = parts[1] if len(parts) > 1 else full_response
            
            current_date = datetime.now().strftime("%B %d, %Y")
            formatted_summary = f"Date: {current_date}\n\n\n{executive_summary}\n\n{detailed_minutes}"
            
            return formatted_summary
        except Exception as e:
            print(f"Summarization error: {e}")
            raise

    def save_summary_to_txt(self, summary, filename="MoM.txt"):
        """Save the summary into a text file"""
        try:
            with open(filename, 'w', encoding='utf-8') as file:
                file.write(summary)
            print(f"Meeting minutes saved to {filename}")
        except IOError as e:
            print(f"Error saving text file: {e}")

    def save_summary_to_docx(self, summary, filename="MoM.docx"):
        """Save the summary into a Word document (.docx)"""
        try:
            doc = Document()
            doc.add_heading('Meeting Minutes', 0)
            doc.add_paragraph(summary)
            doc.save(filename)
            print(f"Meeting minutes saved to {filename}")
        except Exception as e:
            print(f"Error saving DOCX file: {e}")

    def save_summary_to_pdf(self, summary, filename="MoM.pdf"):
        """Save the summary into a PDF document"""
        try:
            c = canvas.Canvas(filename, pagesize=letter)
            c.drawString(100, 750, "Meeting Minutes")
            text = c.beginText(100, 730)
            text.setFont("Helvetica", 10)
            text.setTextOrigin(100, 730)
            text.textLines(summary)
            c.drawText(text)
            c.save()
            print(f"Meeting minutes saved to {filename}")
        except Exception as e:
            print(f"Error saving PDF file: {e}")

    def process_meeting(self, input_file, output_format="txt"):
        """Complete process from extracting audio to generating minutes"""
        # Step 1: Extract audio from video file
        audio_file = self.extract_audio(input_file)

        # Step 2: Transcribe the audio to text
        transcription = self.transcribe_audio(audio_file)
        print("Transcription completed.")

        # Step 3: Generate meeting minutes summary 
        MoM = self.summarize_transcription(transcription)
        print("Meeting minutes generated.\n")

        # Optional cleanup of temporary audio file
        if os.path.exists(audio_file):
            os.remove(audio_file)

        # Step 4: Save the summary to the specified file format
        output_formats = {
            "txt": self.save_summary_to_txt,
            "docx": self.save_summary_to_docx,
            "pdf": self.save_summary_to_pdf
        }

        save_func = output_formats.get(output_format)
        if save_func:
            save_func(MoM)
        else:
            print("Unsupported format. Please choose txt, docx, or pdf.")

        return MoM

def main():
    # Check if video file exists
    input_video_file = "C:/Users/immanuel/Desktop/projekt/Python/Real-Time-Translator/Meeting2.mp4"
    if not os.path.exists(input_video_file):
        print(f"Error: Video file {input_video_file} not found.")
        return

    # Initialize MeetingMinutesGenerator with model size, device, and compute type
    generator = MeetingMinutesGenerator(model_size="base", device="cpu", compute_type="float32")
    
    # Specify the output format ('txt', 'docx', 'pdf')
    output_format = "docx"

    try:
        # Process the meeting
        MoM = generator.process_meeting(input_video_file, output_format=output_format)

        print("MOM Generated successfully")
    
    except Exception as e:
        print(f"An error occurred: {e}")

    #Email part
    csv_path = "fake_participants.csv"  
    attachment_path = "MOM.docx"    
    # Sender email credentials
    sender_email = "immanuelantony2571@gmail.com"
    sender_password = "tmli vesr bpyq hgkm"  
    # Email details
    subject = "Meeting Minutes Document - Final Version"
    body = "Hello,\n\nPlease find the attached document from the meeting.\n\nBest regards,\nImmanuel Anthony"

    # Extract emails from the CSV file
    email_list = extract_emails(csv_path)
    # Send the email to each recipient
    for recipient in email_list:
        send_email(sender_email, sender_password, recipient, subject, body, attachment_path)
        
    print("Successfully Completed")

if __name__ == "__main__":
    main()